#!/usr/bin/env python3
"""
Markdown 보고서를 .docx 로 변환한다.

pandoc 없이 python-docx 만으로 동작한다. 이 저장소의 보고서가 쓰는 문법에
맞춘 것이므로 범용 변환기는 아니다. 지원 범위:

    # ~ ###### 제목        **굵게**  *기울임*  `인라인 코드`
    문단                    [링크 텍스트](주소)
    - / * / 1. 목록 (중첩)  > 인용
    | 표 | (GFM 파이프)     ``` 코드 블록
    ---  가로줄            ![대체글](그림.png)  그림
    ```mermaid``` 블록은 **그림으로 구워 넣는다** (mmdc 가 있을 때).
    ASCII 로 그린 그림은 docx 에서 깨진다 — 고정폭 글꼴이라도 한글과 罫線의
    폭이 맞지 않는다. 그래서 도표는 mermaid 로 쓰고 여기서 굽는다.

사용:
    python md2docx.py docs/보고서.md
    python md2docx.py docs/ -o build/       # 디렉토리 전체
"""
import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 한글이 섞이므로 본문은 한글 글꼴, 코드는 고정폭으로 둔다.
BODY_FONT = "맑은 고딕"
CODE_FONT = "D2Coding"
CODE_FALLBACK = "Consolas"

RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
RE_HR = re.compile(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$")
RE_FENCE = re.compile(r"^\s*```+\s*(\w*)\s*$")
RE_LIST = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$")
RE_QUOTE = re.compile(r"^\s*>\s?(.*)$")
RE_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$")
RE_IMAGE = re.compile(r"^\s*!\[(?P<alt>[^\]]*)\]\((?P<src>[^)\s]+)\)\s*$")

# 용지. ERD 처럼 도표만 있는 문서는 A3 가로로 뽑으면 글자가 세 배로 커진다.
PAPERS = {"a4": (21.0, 29.7), "a3": (29.7, 42.0), "letter": (21.59, 27.94)}
MARGIN_CM = 2.0

# 그림이 쓸 수 있는 폭·높이. 용지에서 여백을 빼고, 거기서 조금 더 뺀다.
# 그림이 이보다 크면 줄인다 — 폭만 맞추면 세로로 긴 도표가 페이지를 넘어가 잘린다.
#
# **높이를 너무 짜게 잡지 않는다.** 세로로 긴 도표는 높이로 갇히는데, 그러면 폭이
# 함께 줄어 글자가 작아진다. A4 에서 21 cm 로 뒀더니 ER 도표가 5.9pt 였다.
MAX_W_CM, MAX_H_CM = 15.5, 23.5


def set_paper(doc, paper: str, landscape: bool):
    """용지를 정하고, 그림이 쓸 수 있는 넓이를 거기 맞춰 다시 잡는다."""
    global MAX_W_CM, MAX_H_CM
    w, h = PAPERS[paper.lower()]
    if landscape:
        w, h = h, w
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(w), Cm(h)
    sec.orientation = (WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT)
    for side in ("left", "right", "top", "bottom"):
        setattr(sec, f"{side}_margin", Cm(MARGIN_CM))
    MAX_W_CM = w - 2 * MARGIN_CM - 0.5
    MAX_H_CM = h - 2 * MARGIN_CM - 0.5

# mermaid 를 그림으로 굽는다. 없으면 코드 블록으로 떨어진다 — 변환이 실패하는
# 것보다 낫다. 굽는 데 헤드리스 크롬이 필요해서 어느 장비에나 있지는 않다.
#
#   npm i @mermaid-js/mermaid-cli && MMDC=$(pwd)/node_modules/.bin/mmdc python md2docx.py …
MMDC = os.environ.get("MMDC") or "mmdc"
# 구운 그림을 두는 곳. 원본이 같으면 다시 굽지 않는다(해시로 가른다).
ASSET_DIR = os.environ.get("MD2DOCX_ASSETS", "assets")

# 인라인: 코드 -> 링크 -> 굵게 -> 기울임 순으로 잘라 낸다.
# 코드를 먼저 처리해야 `**` 가 코드 안에 있을 때 굵게로 오인하지 않는다.
RE_INLINE = re.compile(
    r"(?P<code>`+[^`]+`+)"
    r"|(?P<link>\[(?P<ltext>[^\]]+)\]\((?P<lhref>[^)]+)\))"
    r"|(?P<bold>\*\*[^*]+\*\*)"
    r"|(?P<ital>(?<!\*)\*[^*]+\*(?!\*))"
)


def _set_font(run, name):
    """python-docx 는 동아시아 글꼴을 따로 지정해야 한글에 적용된다."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_inline(par, text, base_font=BODY_FONT, bold=False):
    """인라인 서식을 해석해 run 들을 붙인다."""
    pos = 0
    for m in RE_INLINE.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()])
            r.bold = bold
            _set_font(r, base_font)
        if m.group("code"):
            r = par.add_run(m.group("code").strip("`"))
            _set_font(r, CODE_FONT)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC0, 0x36, 0x2C)
        elif m.group("link"):
            # 하이퍼링크 관계를 만드는 대신 텍스트만 남기고 주소를 괄호로 덧붙인다.
            # 상대경로 링크(../devlog/...)가 많아 클릭 가능한 링크는 의미가 없다.
            r = par.add_run(m.group("ltext"))
            r.bold = bold
            r.underline = True
            _set_font(r, base_font)
        elif m.group("bold"):
            r = par.add_run(m.group("bold").strip("*"))
            r.bold = True
            _set_font(r, base_font)
        elif m.group("ital"):
            r = par.add_run(m.group("ital").strip("*"))
            r.italic = True
            r.bold = bold
            _set_font(r, base_font)
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:])
        r.bold = bold
        _set_font(r, base_font)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=len(rows), cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ""
        add_inline(c.paragraphs[0], cell, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(c, "EDF0F4")
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            c = t.rows[i].cells[j]
            c.text = ""
            add_inline(c.paragraphs[0], row[j] if j < len(row) else "")
    doc.add_paragraph()


def add_code(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(14)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        r = par.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        _set_font(r, CODE_FONT)
        r.font.size = Pt(9)
    _shade_paragraph(par, "F4F5F7")


def add_image(doc, path: Path, alt: str = ""):
    """그림 한 장. 폭과 높이 **둘 다** 본문 안에 들어오게 줄인다."""
    from PIL import Image                                      # noqa: PLC0415
    with Image.open(path) as im:
        w_px, h_px = im.size
    w, h = MAX_W_CM, MAX_W_CM * h_px / w_px
    if h > MAX_H_CM:
        h, w = MAX_H_CM, MAX_H_CM * w_px / h_px
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(path), width=Cm(w), height=Cm(h))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def render_mermaid(src: str, base_dir: Path) -> Path | None:
    """mermaid 원문을 PNG 로. 못 구우면 None 을 낸다.

    **원본이 같으면 다시 굽지 않는다.** 헤드리스 크롬을 띄우는 데 몇 초가 걸려서,
    문서를 한 줄 고칠 때마다 도표를 전부 다시 구우면 변환이 느려진다.
    """
    out_dir = base_dir / ASSET_DIR
    out = out_dir / f"mermaid-{hashlib.sha1(src.encode()).hexdigest()[:12]}.png"
    if out.exists():
        return out
    out_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as td:
        mmd = Path(td) / "d.mmd"
        mmd.write_text(src, encoding="utf-8")
        # 컨테이너·CI 에서는 크롬 샌드박스를 못 쓴다. 우리가 만든 파일만 그리므로
        # 여기서는 꺼도 된다.
        cfg = Path(td) / "p.json"
        cfg.write_text(json.dumps({"args": ["--no-sandbox",
                                            "--disable-setuid-sandbox",
                                            "--disable-dev-shm-usage"]}))
        cmd = [MMDC, "-p", str(cfg), "-i", str(mmd), "-o", str(out),
               "-b", "white", "-s", "3"]     # -s 3: 인쇄해도 글자가 뭉개지지 않게
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            print(f"  mermaid 를 못 구웠다 ({e.__class__.__name__}) — 코드로 남긴다",
                  file=sys.stderr)
            return None
        if r.returncode != 0 or not out.exists():
            print(f"  mermaid 를 못 구웠다: {r.stderr.strip()[:200]}", file=sys.stderr)
            return None
    return out


def _shade_paragraph(par, hex_color):
    ppr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    ppr.append(shd)


def convert(md_path: Path, out_path: Path, paper="a4", landscape=False):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_paper(doc, paper, landscape)

    # 기본 스타일. 한글이 적용되려면 eastAsia 글꼴을 따로 지정해야 한다.
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = BODY_FONT
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # 코드 블록 — mermaid 면 그림으로 굽는다
        m = RE_FENCE.match(line)
        if m:
            lang = (m.group(1) or "").lower()
            buf = []
            i += 1
            while i < n and not RE_FENCE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            i += 1
            png = (render_mermaid("\n".join(buf), md_path.parent)
                   if lang == "mermaid" else None)
            if png:
                add_image(doc, png)
            else:
                add_code(doc, buf)
            continue

        # 그림
        m = RE_IMAGE.match(line)
        if m:
            src = (md_path.parent / m.group("src")).resolve()
            if src.exists():
                add_image(doc, src, m.group("alt"))
            else:
                print(f"  그림이 없다: {m.group('src')}", file=sys.stderr)
            i += 1
            continue

        # 표 — 다음 줄이 구분선이면 표로 본다
        if "|" in line and i + 1 < n and RE_TABLE_SEP.match(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if RE_HR.match(line):
            p = doc.add_paragraph()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "C8CDD5")
            pbdr.append(bottom)
            p._p.get_or_add_pPr().append(pbdr)
            i += 1
            continue

        m = RE_HEADING.match(line)
        if m:
            level = len(m.group(1))
            par = doc.add_heading("", level=min(level, 4))
            add_inline(par, m.group(2), bold=False)
            for r in par.runs:
                r.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)
            i += 1
            continue

        m = RE_QUOTE.match(line)
        if m:
            buf = [m.group(1)]
            i += 1
            while i < n and RE_QUOTE.match(lines[i]):
                buf.append(RE_QUOTE.match(lines[i]).group(1))
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            add_inline(par, " ".join(x for x in buf if x.strip()))
            for r in par.runs:
                r.italic = True
            continue

        m = RE_LIST.match(line)
        if m:
            indent = len(m.group(1))
            ordered = not m.group(2)[0] in "-*+"
            body = m.group(3)
            # 다음 줄이 들여쓴 이어짐이면 합친다
            i += 1
            while i < n and lines[i].strip() and not RE_LIST.match(lines[i]) \
                    and lines[i].startswith(" " * (indent + 2)) \
                    and not RE_HEADING.match(lines[i]):
                body += " " + lines[i].strip()
                i += 1
            style_name = "List Number" if ordered else "List Bullet"
            level = min(indent // 2, 2)
            if level:
                style_name += f" {level + 1}"
            try:
                par = doc.add_paragraph(style=style_name)
            except KeyError:
                par = doc.add_paragraph(style="List Bullet")
            add_inline(par, body)
            continue

        if not line.strip():
            i += 1
            continue

        # 일반 문단 — 빈 줄까지 이어 붙인다 (md 의 줄바꿈은 문단 안에서 무시)
        buf = [line.strip()]
        i += 1
        while i < n and lines[i].strip() and not RE_HEADING.match(lines[i]) \
                and not RE_LIST.match(lines[i]) and not RE_FENCE.match(lines[i]) \
                and not RE_HR.match(lines[i]) and not RE_QUOTE.match(lines[i]) \
                and "|" not in lines[i]:
            buf.append(lines[i].strip())
            i += 1
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(7)
        add_inline(par, " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("target", help=".md 파일 또는 디렉토리")
    ap.add_argument("-o", "--out", help="출력 파일 또는 디렉토리 (기본: 원본과 같은 자리)")
    ap.add_argument("--paper", default="a4", choices=sorted(PAPERS),
                    help="용지 (기본 a4). 도표만 있는 문서는 a3 가 낫다")
    ap.add_argument("--landscape", action="store_true", help="가로로")
    args = ap.parse_args()

    target = Path(args.target)
    if target.is_dir():
        files = sorted(target.rglob("*.md"))
    elif target.exists():
        files = [target]
    else:
        raise SystemExit(f"찾을 수 없다: {target}")
    if not files:
        raise SystemExit(f".md 파일이 없다: {target}")

    out = Path(args.out) if args.out else None
    for f in files:
        if out and out.suffix.lower() == ".docx":
            dest = out
        elif out:
            dest = out / f.with_suffix(".docx").name
        else:
            dest = f.with_suffix(".docx")
        convert(f, dest, args.paper, args.landscape)
        print(f"{f} -> {dest}  ({dest.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
